#!/usr/bin/env python3
"""
The Founders Korea 사업 전략서 - DOCX 생성 스크립트
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime


def set_cell_shading(cell, color_hex):
    """셀 배경색 설정"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_table_borders(table):
    """테이블 테두리 설정"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'CCCCCC')
        borders.append(border)
    tblPr.append(borders)


def create_styled_table(doc, headers, rows, header_color='1B365D'):
    """스타일된 테이블 생성"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)

    # 헤더
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, header_color)

    # 데이터
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'F5F7FA')

    return table


def main():
    doc = Document()

    # ─── 기본 스타일 설정 ───
    style = doc.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(6)

    # Heading 스타일
    for level in range(1, 4):
        h_style = doc.styles[f'Heading {level}']
        h_style.font.name = 'Malgun Gothic'
        h_style.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

    # ─── 표지 ───
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run('The Founders Korea')
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('사업 전략서')
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x0D, 0x73, 0x77)

    doc.add_paragraph()

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run('창업가를 위한 실전 미디어 플랫폼')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph()
    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('한국창업경영센터\n')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    run = info.add_run(f'\n{datetime.date.today().strftime("%Y년 %m월")}')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version.add_run('v3 - 의회 최종 검토 반영')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xD4, 0xA8, 0x53)
    run.bold = True

    doc.add_page_break()

    # ─── 목차 ───
    doc.add_heading('목차', level=1)
    toc_items = [
        '1. Executive Summary',
        '2. 사업 배경 및 시장 환경',
        '3. 사업 목적 (4대 축)',
        '4. 타겟 고객 세분화',
        '5. 콘텐츠 전략 (5대 카테고리)',
        '6. 킬러 피처: 정부지원사업 DB (Gov-Support Radar)',
        '7. AI 맞춤 추천 시스템',
        '8. 리드 퍼널 전략',
        '9. 차별화 전략',
        '10. 수익 모델 (Freemium)',
        '11. 멀티미디어 전략',
        '12. GTM (Go-to-Market) 전략',
        '13. 성장 전략 및 로드맵',
        '14. 기술 아키텍처',
        '15. KPI 및 성과 지표',
        '16. 리스크 관리',
        '17. 비용 구조',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ─── 1. Executive Summary ───
    doc.add_heading('1. Executive Summary', level=1)

    p = doc.add_paragraph()
    run = p.add_run('The Founders Korea')
    run.bold = True
    p.add_run('는 한국창업경영센터가 운영하는 ')
    run = p.add_run('실전 창업 미디어 플랫폼')
    run.bold = True
    p.add_run('입니다. 기존 창업 미디어가 투자 유치와 유니콘 기업 소식에 편중된 반면, '
              '본 플랫폼은 소상공인과 초기 스타트업 대표에게 필요한 경영 실무 정보, '
              '정부지원사업 데이터베이스, AI 맞춤 추천, 시민 참여 저널리즘을 결합한 '
              '차별화된 서비스를 제공합니다.')

    doc.add_heading('핵심 가치 제안', level=2)
    bullets = [
        ('508개+ 정부지원사업 구조화 DB', '검색, 필터, AI 매칭으로 약 9조원 규모의 지원 생태계를 한 곳에서 접근'),
        ('5대 카테고리 콘텐츠', 'Growth Playbook, AI/DX Lab, Founder\'s Voice, Gov-Support Radar, Success Cases'),
        ('AI 맞춤 추천 시스템', '온보딩 퀴즈를 통한 개인화된 정부지원사업 매칭'),
        ('콘텐츠→컨설팅 원스톱', '정보 소비에서 전문가 상담까지 하나의 플랫폼에서 연결'),
        ('Freemium 전략', '6개월간 완전 무료 → 데이터 기반 유료화 결정'),
    ]
    for title, desc in bullets:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(title + ': ')
        run.bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ─── 2. 사업 배경 및 시장 환경 ───
    doc.add_heading('2. 사업 배경 및 시장 환경', level=1)

    doc.add_heading('시장 문제', level=2)
    problems = [
        '기존 창업 미디어는 투자 유치, 유니콘 기업 등 화려한 소식에 편중',
        '소상공인/초기 스타트업 대표에게 필요한 경영 실무 정보 부재',
        '세무, 법률, 정부지원사업 등 당장 필요한 정보를 찾기 어려움',
        '창업가 간 현장 경험을 공유할 수 있는 양방향 미디어 플랫폼 부재',
        '정부지원사업 정보가 분산되어 접근성이 낮음 (508개 사업, 111개 기관)',
    ]
    for prob in problems:
        doc.add_paragraph(prob, style='List Bullet')

    doc.add_heading('2026년 시장 환경', level=2)
    create_styled_table(doc,
        ['지표', '수치', '시사점'],
        [
            ['정부 창업지원 예산', '3.4조원 (역대 최대)', '정부지원 정보 수요 폭증'],
            ['소상공인 지원 예산', '5.4조원 (역대 최대)', '소상공인 타겟 가치 상승'],
            ['소상공인 AI 도입 효과', '82% 생산성 향상', 'AI/DX 콘텐츠 필수'],
            ['정부 정책 방향', '예비창업 → 사업화/글로벌', '성장 단계별 콘텐츠 필요'],
            ['콘텐츠 트렌드', '쇼트폼/영상 주류화', '텍스트 전용 전략 한계'],
            ['AI 콘텐츠 자동화', '생성 콘텐츠 80% 도달 예상', 'AI 활용 운영 효율화'],
        ]
    )
    doc.add_paragraph()

    doc.add_heading('시장 기회', level=2)
    opportunities = [
        '소상공인 약 700만+ (자영업자 포함)',
        '연간 신규 창업 약 100만+ 건',
        '정부 창업 지원사업 508개, 정보 접근 허들 높음 → 구조화 DB의 기회',
        '"창업" "소상공인" "정부지원사업" 검색량 지속 증가',
        '기업가형 소상공인 육성 패러다임 → 성장 지향 콘텐츠 수요',
    ]
    for opp in opportunities:
        doc.add_paragraph(opp, style='List Bullet')

    doc.add_page_break()

    # ─── 3. 사업 목적 ───
    doc.add_heading('3. 사업 목적 (4대 축)', level=1)

    create_styled_table(doc,
        ['축', '명칭', '설명'],
        [
            ['1', '실전 경영 + AI/DX 인사이트', '현장에서 바로 쓸 수 있는 정보와 AI 활용법 제공'],
            ['2', '시민 참여 저널리즘', '창업가가 직접 투고하는 현장 중심 콘텐츠'],
            ['3', '실용 도구 제공', '정부지원사업 DB, AI 도구 가이드 등 도구형 자산'],
            ['4', '공신력 확보 & 리드 창출', '한국창업경영센터 컨설팅 전문성의 미디어화'],
        ],
        header_color='0D7377'
    )

    doc.add_page_break()

    # ─── 4. 타겟 고객 세분화 ───
    doc.add_heading('4. 타겟 고객 세분화', level=1)

    doc.add_heading('Primary: 현직 사업자', level=2)
    create_styled_table(doc,
        ['세분화', '프로필', '핵심 니즈', '접근 채널'],
        [
            ['소상공인', '음식점, 카페, 소매업 등', '세무, AI 도입, 배달앱, 인건비 관리', '네이버, 카카오'],
            ['초기 스타트업', '팀 5명 이하, 시드~프리A', '법인 운영, 정부지원사업, 투자', '구글, 커뮤니티'],
            ['프랜차이즈 예비 가맹주', '창업 자금 준비 완료', '가맹 계약, 입지 분석, 수익 구조', '네이버 카페, 유튜브'],
        ]
    )
    doc.add_paragraph()

    doc.add_heading('Secondary: 예비 창업자', level=2)
    create_styled_table(doc,
        ['세분화', '프로필', '핵심 니즈'],
        [
            ['퇴직 후 창업 준비', '40~50대', '사업 아이템, 리스크 관리, 정부지원'],
            ['청년 창업', '20~30대', '정부지원사업, 사업계획서, AI 활용'],
            ['부업/사이드 프로젝트', '직장인', '사업자 등록, 세금, 온라인 사업'],
        ]
    )
    doc.add_paragraph()

    doc.add_heading('Tertiary: 생태계 관계자', level=2)
    create_styled_table(doc,
        ['세분화', '프로필', '활용 목적'],
        [
            ['투자자', '엔젤, VC', '시장 동향, 딜소싱'],
            ['컨설턴트/전문가', '세무사, 변호사', '고객 확보, 전문성 노출'],
            ['정책 담당자', '중기부, 지자체', '현장 목소리 파악'],
        ]
    )
    doc.add_paragraph()

    doc.add_heading('초기 집중 페르소나 (하이퍼 타겟팅)', level=2)
    p = doc.add_paragraph()
    run = p.add_run('1순위 페르소나: ')
    run.bold = True
    p.add_run('B2B SaaS 파운더, 창업 3년 차, 5~10명 규모, 스케일업용 R&D 정부지원사업 탐색 중')
    p = doc.add_paragraph()
    p.add_run('이 페르소나 기준으로 콘텐츠/DB/UX 설계 → 검증 후 다른 페르소나로 확장.')

    doc.add_page_break()

    # ─── 5. 콘텐츠 전략 (5대 카테고리) ───
    doc.add_heading('5. 콘텐츠 전략 (5대 카테고리)', level=1)

    p = doc.add_paragraph()
    run = p.add_run('우선순위: ')
    run.bold = True
    p.add_run('Gov-Support Radar > Growth Playbook > Success Cases > AI/DX Lab > Founder\'s Voice')

    doc.add_heading('카테고리별 상세', level=2)
    create_styled_table(doc,
        ['카테고리', '작성자', '콘텐츠 유형', '발행 빈도', 'KPI'],
        [
            ['Growth Playbook', '센터 전문가', '세무/법률/경영 심층 리포트', '주 1~2회', '검색 유입, 체류 시간'],
            ['AI/DX Lab', '센터 + 외부 전문가', 'AI 도구 비교, 도입 가이드, 사례', '주 1회', '도구 클릭률, 공유 수'],
            ['Founder\'s Voice', '창업가 직접 투고', '현장 노하우, 실패/성공기', '주 2~3회', '투고 수, 재방문율'],
            ['Gov-Support Radar', '스크래퍼 + 편집자', '정부지원 DB + 주간 뉴스', 'DB 상시, 뉴스 주 1회', 'DB 검색 수, 구독 전환'],
            ['Success Cases', '한국창업경영센터', '컨설팅 성공 사례', '월 2~4회', '상담 신청 전환율'],
        ]
    )
    doc.add_paragraph()

    doc.add_heading('5대 카테고리 전략적 플라이휠', level=2)
    flywheel_items = [
        ('Growth Playbook', '전문성 증명 → SEO 트래픽 + 신뢰 구축'),
        ('AI/DX Lab', '실전 도구 → AI 도입 니즈 충족 + 제휴 수익'),
        ('Founder\'s Voice', '커뮤니티 → UGC 확장 + 충성도'),
        ('Gov-Support Radar', '실용 DB → 고빈도 방문 + 구독자 확보 (킬러 피처)'),
        ('Success Cases', '직접 전환 → 상담 리드 창출'),
    ]
    for cat, desc in flywheel_items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(cat + ': ')
        run.bold = True
        p.add_run(desc)

    p = doc.add_paragraph()
    run = p.add_run('전체 연결: ')
    run.bold = True
    p.add_run('뉴스레터 구독 → 상담 신청 → 컨설팅 계약')

    doc.add_heading('카테고리별 비즈니스 연결', level=2)
    create_styled_table(doc,
        ['카테고리', '예시 콘텐츠', 'CTA', '전환 경로'],
        [
            ['Growth Playbook', '2026년 소상공인 세무 가이드', '세무 전문가에게 상담받기', '/contact'],
            ['AI/DX Lab', '소상공인 AI 고객응대 툴 TOP 5', '도구 제휴 링크 + AI 도입 컨설팅', '/contact'],
            ['Founder\'s Voice', '배달앱 수수료와 싸우며 배운 것들', '나도 이야기 들려주기', '/submit'],
            ['Gov-Support Radar', 'AI 도입 지원 검색', '이 사업 신청 도움받기', '/contact (가장 직접적 전환)'],
            ['Success Cases', '매출 3억 음식점 법인 전환 사례', '비슷한 고민? 무료 상담', '/contact'],
        ]
    )

    doc.add_page_break()

    # ─── 6. 킬러 피처: 정부지원사업 DB ───
    doc.add_heading('6. 킬러 피처: 정부지원사업 DB (Gov-Support Radar)', level=1)

    doc.add_heading('왜 킬러 피처인가?', level=2)
    reasons = [
        '타겟 70%가 "정부지원 정보 접근이 어렵다"고 응답',
        '508개 사업, 111개 기관에 분산된 정보',
        '연간 3.4조 + 5.4조 = 약 9조원 규모의 지원 생태계',
        '구조화된 검색 DB = 압도적 실용 가치 + SEO 트래픽 엔진',
    ]
    for r in reasons:
        doc.add_paragraph(r, style='List Bullet')

    doc.add_heading('DB 데이터 필드', level=2)
    fields = [
        '사업명, 주관기관, 신청기간(시작~마감), 지원대상(업종/지역/규모)',
        '지원내용(금액/유형), 신청방법, 원문링크, 태그',
        '상태(예정/접수중/마감), 수집일시',
    ]
    for f in fields:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_heading('데이터 수집 전략', level=2)
    create_styled_table(doc,
        ['소스', '수집 방법', '우선순위'],
        [
            ['K-Startup (k-startup.go.kr)', 'RSS + 스크래핑', '높음'],
            ['중소벤처기업부 (mss.go.kr)', '공고 페이지 스크래핑', '높음'],
            ['공공데이터포털 API', 'REST API', '높음'],
            ['창업진흥원 (kised.or.kr)', '스크래핑', '중간'],
            ['지자체 창업 지원', '주요 시도 스크래핑', '중간'],
            ['플래텀, 벤처스퀘어', 'RSS 피드', '낮음'],
        ]
    )
    doc.add_paragraph()

    doc.add_heading('커뮤니티 리뷰 시스템', level=2)
    create_styled_table(doc,
        ['기능', '설명', '도입 시기'],
        [
            ['성공률 통계', '프로그램별 실제 선정율, 평균 지원 규모', 'Phase 2'],
            ['파운더 리뷰', '실제 신청자/수혜자의 팁, 실패 사례, 주의사항', 'Phase 2'],
            ['Q&A 게시판', '프로그램별 질문/답변 커뮤니티', 'Phase 2'],
            ['Founder\'s Pick', '주간 Top 추천 프로그램 큐레이션', 'Phase 1'],
        ]
    )
    doc.add_paragraph()

    doc.add_heading('신청 키트 (Application Kits)', level=2)
    kits = [
        '템플릿: 주요 프로그램별 사업계획서 템플릿',
        '체크리스트: 필수 서류, 신청 절차 단계별 가이드',
        'Step-by-step 가이드: 인기 프로그램 Top 20 상세 신청 가이드',
    ]
    for k in kits:
        doc.add_paragraph(k, style='List Bullet')

    doc.add_heading('vs 기존 정부 포털 차별화', level=2)
    create_styled_table(doc,
        ['항목', 'bizinfo.go.kr', 'K-Startup', 'The Founders Korea'],
        [
            ['데이터', '원본 데이터 나열', '정부 공식 포털', '큐레이션 + AI 매칭'],
            ['UX', '관료적 UI', '기능 복잡', '직관적 검색/필터'],
            ['맞춤 추천', '없음', '기본 필터', 'AI 매칭 점수 + 퀴즈'],
            ['커뮤니티', '없음', '없음', '파운더 리뷰 + Q&A'],
            ['실행 지원', '없음', '일부 교육', '신청 키트 + 컨설팅'],
            ['포지셔닝', '정보 제공', '정보 제공', 'Actionable Intelligence'],
        ]
    )

    doc.add_page_break()

    # ─── 7. AI 맞춤 추천 시스템 ───
    doc.add_heading('7. AI 맞춤 추천 시스템', level=1)

    doc.add_heading('온보딩 퀴즈 Flow', level=2)
    p = doc.add_paragraph()
    p.add_run('사용자가 간단한 퀴즈(4~5개 질문)를 통해 프로필을 입력하면, '
              'AI 매칭 엔진이 적합한 정부지원사업을 점수 기반으로 추천합니다.')

    doc.add_heading('퀴즈 항목', level=3)
    quiz_items = [
        '업종 (예: B2B SaaS, 음식점, 제조업 등)',
        '사업 단계 (예비창업/초기/성장/안정)',
        '규모 (1인/5인 이하/10인 이하/10인 이상)',
        '지역 (시/도)',
        '관심 분야 (R&D, 자금, 교육, 공간, 멘토링)',
    ]
    for item in quiz_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('AI 매칭 엔진 가중치', level=2)
    create_styled_table(doc,
        ['매칭 요소', '배점', '매칭 로직'],
        [
            ['지역', '25점', '정확 일치 = 1.0, 전국 프로그램 = 0.8'],
            ['업종', '25점', '매칭 업종 세트 기준'],
            ['기업 규모', '20점', '정확 일치 = 1.0'],
            ['관심 분야 vs 키워드', '20점', '오버랩 비율'],
            ['프로그램 상태', '10점', '모집중 = 1.0, 상시 = 0.8, 마감 = 0.0'],
        ]
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('결과: ')
    run.bold = True
    p.add_run('매칭 점수(0~100%) 기반 정렬된 추천 프로그램 리스트 + 매칭 이유 설명')

    doc.add_page_break()

    # ─── 8. 리드 퍼널 전략 ───
    doc.add_heading('8. 리드 퍼널 전략', level=1)

    p = doc.add_paragraph()
    run = p.add_run('핵심 취약점: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    p.add_run('"신뢰 → 상담 신청" 전환 단계. 콘텐츠로 유입은 되지만, 상담까지 연결하는 동기 부여가 부족.')

    doc.add_heading('강화된 5단계 퍼널', level=2)
    create_styled_table(doc,
        ['단계', '명칭', '핵심 전략', '주요 액션'],
        [
            ['1', '인지 (Awareness)', 'SEO + 정부지원 DB 검색 유입', 'Growth Playbook 롱테일 키워드, DB 검색 트래픽'],
            ['2', '관심 (Interest)', '콘텐츠 소비 + DB 반복 활용', '다수 페이지뷰, 정부지원 DB 고빈도 방문'],
            ['3', '구독 (Subscribe)', '뉴스레터 + 맞춤 알림', '주간 뉴스 + 정부지원 알림, 이메일 수집'],
            ['4', '전환 (Conversion)', '4개 전환 경로 운영', 'DB→상담, Playbook→상담, Cases→상담, AI매칭→상담'],
            ['5', '계약 (Contract)', '유료 컨설팅 연결', '한국창업경영센터 세무/법률/경영/정부지원 대행'],
        ]
    )
    doc.add_paragraph()

    doc.add_heading('퍼널 전환율 목표', level=2)
    create_styled_table(doc,
        ['단계', '목표', '핵심 전략'],
        [
            ['방문 → 재방문', '35%', '정부지원 DB의 반복 활용 유도'],
            ['방문 → 구독', '5~8%', '맞춤 알림으로 구독 동기 강화'],
            ['구독 → 상담 신청', '8~12%', '무료 콘텐츠로 신뢰 구축 후 자연 전환'],
            ['상담 → 계약', '20~30%', '센터 영업 역량'],
        ]
    )

    doc.add_page_break()

    # ─── 9. 차별화 전략 ───
    doc.add_heading('9. 차별화 전략', level=1)

    doc.add_heading('vs 경쟁 미디어', level=2)
    create_styled_table(doc,
        ['항목', '플래텀/벤처스퀘어', 'The Founders Korea'],
        [
            ['타겟', '투자 유치, 유니콘', '소상공인, 초기 스타트업 실무'],
            ['콘텐츠', '기자 작성 뉴스', '창업가 직접 투고 + 전문가 리포트'],
            ['정부지원 정보', '단순 뉴스 보도', '구조화 DB + AI 매칭 + 커뮤니티 리뷰'],
            ['AI/DX', '산업 동향 기사', '실전 도구 가이드 + 도입 사례'],
            ['수익 모델', '광고', 'Freemium + 컨설팅 리드'],
            ['상호작용', '일방향', '시민기자 투고 + 커뮤니티 리뷰'],
            ['운영 주체', '독립 미디어', '전문 컨설팅 기관 (원스톱)'],
        ]
    )
    doc.add_paragraph()

    doc.add_heading('5대 차별점', level=2)
    diffs = [
        ('정부지원사업 DB', '508개 사업을 검색/필터/알림으로 제공 (타 미디어 없음)'),
        ('AI/DX 실전 가이드', '"이 도구를 이렇게 써라" 수준의 실행 가능한 콘텐츠'),
        ('시민 참여', 'Founder\'s Voice로 현장 진정성 확보 (모방 어려움)'),
        ('콘텐츠 → 상담 원스톱', '정보 소비 → 전문가 연결이 한 플랫폼에서'),
        ('AI 기반 운영', '뉴스 수집, 콘텐츠 초안, 검수 보조 등 AI 활용 효율화'),
    ]
    for i, (title, desc) in enumerate(diffs, 1):
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'{title}: ')
        run.bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ─── 10. 수익 모델 ───
    doc.add_heading('10. 수익 모델 (Full Freemium)', level=1)

    doc.add_heading('수익 전략: 완전 Freemium → 데이터 기반 유료화 결정', level=2)
    p = doc.add_paragraph()
    p.add_run('Phase 1-2 (0~6개월): 모든 콘텐츠/DB/AI 매칭/신청 키트 완전 무료 + 컨설팅 리드 상시 수익화')
    p = doc.add_paragraph()
    p.add_run('Phase 3 (6개월+): 축적된 사용자 데이터를 기반으로 최적 유료화 모델 선택')

    doc.add_heading('Phase별 수익 전략', level=2)
    create_styled_table(doc,
        ['Phase', '기간', '수익 전략', '상세'],
        [
            ['Phase 1', '0~3개월', '컨설팅 리드만', '모든 콘텐츠/DB 무료, 상담 연결만 수익화'],
            ['Phase 2', '3~6개월', '컨설팅 리드 + 제휴', 'AI 도구 어필리에이트 추가, 정부지원 신청 대행 연결'],
            ['Phase 3', '6개월+', '데이터 기반 결정', '축적된 사용자 데이터로 최적 유료화 모델 선택'],
        ]
    )
    doc.add_paragraph()

    doc.add_heading('Phase 3 유료화 옵션', level=2)
    create_styled_table(doc,
        ['옵션', '모델', '적합 조건'],
        [
            ['A) 구독 모델', '프리미엄 DB 기능 + AI 매칭 + 맞춤 알림 월 구독', 'DB 반복 사용률 높을 때'],
            ['B) Value Ladder', '디지털 상품(₩29K) → 리포트(₩99K) → 컨설팅(₩500K+)', '단품 구매 의향 높을 때'],
            ['C) 하이브리드', '기본 구독 + 프리미엄 단품', '둘 다 수요가 있을 때'],
        ]
    )
    doc.add_paragraph()

    doc.add_heading('수익원 상세', level=2)
    create_styled_table(doc,
        ['수익원', '설명', '도입 시기', '예상 가격'],
        [
            ['컨설팅 리드 (핵심)', '콘텐츠 → 상담 → 유료 컨설팅', 'Phase 1 (상시)', '건당 50만~500만+'],
            ['정부지원 신청 대행', 'DB → 신청 컨설팅 연결', 'Phase 2', '건당 50만~200만'],
            ['AI 도구 제휴', 'AI/DX Lab 도구 어필리에이트', 'Phase 2', '건당 수수료'],
            ['디지털 상품', '템플릿, 체크리스트, 가이드북', 'Phase 3', '2.9만~9.9만'],
            ['프리미엄 리포트', '심화 분석 보고서 유료 판매', 'Phase 3', '9.9만~29.9만'],
            ['프리미엄 구독', 'DB 고급 기능 + 맞춤 알림', 'Phase 3', '월 5만'],
            ['교육/세미나', '온라인 강의, 워크숍', 'Phase 3+', '5만~30만'],
        ]
    )

    doc.add_page_break()

    # ─── 11. 멀티미디어 전략 ───
    doc.add_heading('11. 멀티미디어 전략', level=1)

    doc.add_heading('단계적 접근: Crawl → Walk → Run', level=2)
    create_styled_table(doc,
        ['단계', '시기', '형식', '방법'],
        [
            ['Crawl', 'Phase 1', '텍스트 + 이미지 카드', '블로그 포스트 + Canva 소셜 카드'],
            ['Walk', 'Phase 2', '+ 쇼트폼 영상 (핵심)', '인기 글 → 30초~1분 요약 영상, 릴스/쇼츠'],
            ['Run', 'Phase 3', '+ 유튜브/팟캐스트', '전문가 인터뷰, Founder\'s Voice Vlog'],
        ]
    )
    doc.add_paragraph()

    doc.add_heading('쇼트폼 영상 전략 (Phase 2+)', level=2)
    create_styled_table(doc,
        ['콘텐츠 유형', '길이', '플랫폼', '빈도'],
        [
            ['정부지원사업 1분 요약', '30초~1분', '인스타 릴스, 유튜브 쇼츠, 틱톡', '주 1~2회'],
            ['AI 도구 빠른 소개', '30초', '인스타 릴스, 유튜브 쇼츠', '주 1회'],
            ['파운더 인터뷰 클립', '1분', '유튜브 쇼츠', '월 2회'],
            ['경영 팁 카드뉴스 영상', '15~30초', '인스타 릴스', '주 1회'],
        ]
    )
    doc.add_paragraph()

    doc.add_heading('One Source Multi-Use 원칙', level=2)
    osmu = [
        '블로그 포스트 (원본)',
        '뉴스레터 요약',
        '소셜 이미지 카드 2~3개 (카카오, 인스타)',
        '쇼트폼 영상 1개 (30초~1분, Phase 2+)',
        'AI 챗봇 추천 데이터 연동 (Phase 2+)',
        '정부지원 DB 연계 (관련 지원사업 링크)',
    ]
    for item in osmu:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ─── 12. GTM 전략 ───
    doc.add_heading('12. GTM (Go-to-Market) 전략', level=1)

    p = doc.add_paragraph()
    run = p.add_run('핵심 원칙: ')
    run.bold = True
    p.add_run('뉴스레터를 초기 핵심 상품으로, Content SEO + 정부지원 DB 검색 트래픽을 성장 엔진으로 활용.')

    doc.add_heading('1) Content SEO 전략', level=2)
    create_styled_table(doc,
        ['키워드 유형', '예시', '담당 카테고리', '기대 효과'],
        [
            ['정보성 롱테일', '소상공인 정부지원사업 신청 방법', 'Gov-Support Radar', '검색 유입 핵심'],
            ['문제 해결형', '법인 전환 절차 2026', 'Growth Playbook', '전문성 입증'],
            ['AI/DX 실전', '소상공인 AI 도입 비용', 'AI/DX Lab', '트렌드 트래픽'],
            ['비교/리뷰', '정부지원사업 추천 TOP 10', 'Gov-Support Radar', '리드 매그넷 연계'],
        ]
    )
    doc.add_paragraph()

    doc.add_heading('2) 파트너십 전략', level=2)
    create_styled_table(doc,
        ['파트너 유형', '대상', '협업 방식', '우선순위'],
        [
            ['액셀러레이터', '스파크랩, 프라이머 등', '콘텐츠 공유, 포트폴리오사 투고', '높음'],
            ['코워킹 스페이스', '위워크, 패스트파이브', '오프라인 홍보, 세미나', '높음'],
            ['대학 창업센터', '주요 대학 창업지원단', '학생 창업자 대상 콘텐츠', '중간'],
            ['정부기관', '창업진흥원, 소진공', '공식 정보 연계, 콘텐츠 제휴', '중간'],
        ]
    )
    doc.add_paragraph()

    doc.add_heading('3) 커뮤니티 참여 전략', level=2)
    create_styled_table(doc,
        ['커뮤니티', '전략', '주의사항'],
        [
            ['Disquiet', '제품 빌딩 과정 공유, 피드백 요청', '자연스러운 참여, 홍보 자제'],
            ['네이버 카페', '정부지원사업 정보 공유, 질문 답변', '도움이 되는 답변 위주, 링크 스팸 금지'],
            ['카카오 오픈채팅', '창업자 그룹 참여, 정보 공유', '신뢰 구축 후 자연 유입'],
        ]
    )
    doc.add_paragraph()

    doc.add_heading('4) 뉴스레터 퍼스트 전략', level=2)
    newsletter_items = [
        '주간 정부지원사업 업데이트 (신규/마감 임박)',
        '이주의 Growth Playbook 요약',
        'Founder\'s Pick: 추천 지원사업 TOP 3',
        'AI/DX 팁 1개',
        'CTA: "맞춤 상담 신청" 또는 "DB에서 더 찾아보기"',
    ]
    for item in newsletter_items:
        doc.add_paragraph(item, style='List Bullet')

    p = doc.add_paragraph()
    run = p.add_run('뉴스레터 구독자 = 가장 가치 있는 리드.')
    run.bold = True
    p.add_run(' 이메일 수집을 모든 페이지의 핵심 CTA로 설정.')

    doc.add_heading('5) 리드 매그넷 전략', level=2)
    create_styled_table(doc,
        ['리드 매그넷', '형식', '대상 페르소나', 'CTA'],
        [
            ['"내 업종별 정부지원사업 TOP 10" 체크리스트', 'PDF', '전체', '이메일 입력 → 다운로드'],
            ['"사업계획서 작성 가이드 + 템플릿"', 'PDF + Excel', '예비 창업자', '이메일 입력 → 다운로드'],
            ['"AI 도입 ROI 계산기"', '웹 도구', '소상공인', '결과 이메일 발송'],
            ['"정부지원사업 신청 체크리스트"', 'PDF', '전체', '이메일 입력 → 다운로드'],
        ]
    )

    doc.add_page_break()

    # ─── 13. 성장 전략 및 로드맵 ───
    doc.add_heading('13. 성장 전략 및 로드맵', level=1)

    doc.add_heading('Phase별 성장 목표', level=2)
    create_styled_table(doc,
        ['Phase', '기간', '핵심 목표', '콘텐츠', '트래픽', '전환'],
        [
            ['1', '0~3개월', '플랫폼 + DB MVP + 뉴스레터', '50개', '월 3,000', '구독 200, 상담 월 5'],
            ['2', '3~6개월', 'CMS + AI 매칭 + 쇼트폼', '100개', '월 10,000', '구독 1,000, 상담 월 15'],
            ['3', '6~12개월', '데이터 기반 유료화 결정', '200개', '월 30,000', '구독 3,000, 상담 월 30'],
            ['이후', '12개월+', '프리미엄 + 커뮤니티 확장', '400+', '월 100,000', '구독 10,000+'],
        ]
    )
    doc.add_paragraph()

    doc.add_heading('트래픽 확보 전략', level=2)
    create_styled_table(doc,
        ['채널', '전략', '기대 비중'],
        [
            ['SEO', '롱테일 키워드 (세무, 정부지원, AI 도입)', '50%'],
            ['정부지원 DB', '"정부지원사업 검색" 직접 유입', '20%'],
            ['뉴스레터', 'Weekly 뉴스 + 맞춤 알림 바이럴', '10%'],
            ['소셜 미디어', '카카오 채널, 인스타, 네이버 블로그', '10%'],
            ['Founder\'s Voice', '투고자 자발적 공유', '5%'],
            ['제휴/커뮤니티', '창업 카페, 파트너 채널', '5%'],
        ]
    )
    doc.add_paragraph()

    doc.add_heading('개발 로드맵', level=2)

    doc.add_heading('Phase 1: 미디어 플랫폼 + DB MVP (0~3개월)', level=3)
    phase1_items = [
        '프로젝트 구조 리팩토링 (app/ 디렉토리 모듈화)',
        'DB 도입 (SQLAlchemy + Alembic, Post + GovSupportProgram 모델)',
        '5대 카테고리 라우트 구현',
        '정부지원사업 DB MVP (검색, 필터, 상세보기, 공공데이터포털 API)',
        'AI 맞춤 추천 온보딩 퀴즈',
        '커뮤니티 리뷰 시스템',
        '신청 키트 MVP',
        '메인 페이지 개편 + SEO 메타 태그',
        'RSS 피드 + 페이지네이션',
    ]
    for item in phase1_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Phase 2: CMS + 투고 + AI 챗봇 (3~6개월)', level=3)
    phase2_items = [
        '관리자 인증 + CMS 대시보드',
        'Founder\'s Voice 투고 폼 + 검토 워크플로',
        'Gov-Support Radar 스크래퍼 확장 (다수 소스 자동 수집)',
        '정부지원 맞춤 알림 + 뉴스레터 구독/발송',
        'AI 챗봇 추천 MVP',
        '통합 검색 기능 (PostgreSQL Full-Text Search)',
        '태그 시스템 + PostgreSQL 전환',
    ]
    for item in phase2_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Phase 3: 데이터 기반 수익화 (6~12개월)', level=3)
    phase3_items = [
        '상담 신청 페이지 (/contact)',
        '6개월 데이터 분석 → 최적 수익 모델 결정',
        '디지털 상품 시스템 (결제 연동)',
        'Success Cases CTA 강화',
        'AI 도구 제휴/어필리에이트',
        '프리미엄 리포트/교육 + 구독 서비스',
        '분석 대시보드 (Google Analytics 연동)',
    ]
    for item in phase3_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ─── 14. 기술 아키텍처 ───
    doc.add_heading('14. 기술 아키텍처', level=1)

    doc.add_heading('기술 스택', level=2)
    create_styled_table(doc,
        ['영역', '기술', '비고'],
        [
            ['백엔드 프레임워크', 'Python FastAPI + Uvicorn', '비동기 고성능'],
            ['ORM', 'SQLAlchemy 2.0 + Alembic', '마이그레이션 관리'],
            ['데이터베이스', 'SQLite (개발) / PostgreSQL (프로덕션)', 'Railway 배포'],
            ['템플릿 엔진', 'Jinja2', '서버 사이드 렌더링'],
            ['CSS 프레임워크', 'TailwindCSS 기반 커스텀', 'Trust Blue 디자인 시스템'],
            ['동적 UI', 'HTMX', '경량 SPA-like 인터랙션'],
            ['마크다운', 'Python-Markdown + Pygments', '코드 하이라이팅 지원'],
            ['배포', 'Railway + Docker', '자동 배포 (git push)'],
            ['AI API', 'OpenAI / Claude API', '매칭, 챗봇, 콘텐츠 보조'],
        ]
    )
    doc.add_paragraph()

    doc.add_heading('데이터베이스 스키마', level=2)
    create_styled_table(doc,
        ['테이블', '주요 필드', '용도'],
        [
            ['categories', 'id, name, slug, description, icon, color', '5대 카테고리 관리'],
            ['posts', 'id, title, slug, content_md/html, category_id, author', '통합 콘텐츠 관리'],
            ['gov_support_programs', 'id, title, organization, support_type, region, deadline, status', '정부지원사업 DB'],
            ['program_reviews', 'id, program_id, nickname, rating, content, success_tag', '커뮤니티 리뷰'],
            ['application_kits', 'id, program_id, checklist_md/html, guide_md/html', '신청 키트'],
            ['user_profiles', 'id, session_id, industry, business_stage, region, interests', 'AI 매칭용 프로필'],
        ]
    )
    doc.add_paragraph()

    doc.add_heading('디자인 시스템', level=2)
    create_styled_table(doc,
        ['요소', '값', '용도'],
        [
            ['Primary Color', '#1B365D (Deep Trust Blue)', '안정감, 전문성'],
            ['Secondary Color', '#0D7377 (Growth Teal)', '액션, 성장'],
            ['Accent Color', '#D4A853 (Warm Gold)', '프리미엄, 주목'],
            ['폰트', 'Pretendard Variable + Noto Sans KR', '한국어 최적화'],
            ['폰트 굵기', '400, 500, 700, 900', '계층적 타이포그래피'],
        ]
    )

    doc.add_page_break()

    # ─── 15. KPI 및 성과 지표 ───
    doc.add_heading('15. KPI 및 성과 지표', level=1)

    doc.add_heading('콘텐츠 지표', level=2)
    create_styled_table(doc,
        ['지표', '3개월', '6개월', '12개월'],
        [
            ['발행 포스트 수', '50', '100', '200'],
            ['정부지원 DB 등록 건수', '200', '400', '500+'],
            ['시민기자 투고 수', '10', '30', '80'],
            ['AI/DX Lab 가이드 수', '10', '25', '50'],
        ]
    )
    doc.add_paragraph()

    doc.add_heading('트래픽 지표', level=2)
    create_styled_table(doc,
        ['지표', '3개월', '6개월', '12개월'],
        [
            ['월간 방문자', '3,000', '10,000', '30,000'],
            ['정부지원 DB 검색 수', '500', '3,000', '10,000'],
            ['페이지뷰/방문', '2.0', '2.5', '3.0'],
            ['평균 체류 시간', '2분', '3분', '4분'],
        ]
    )
    doc.add_paragraph()

    doc.add_heading('전환 지표', level=2)
    create_styled_table(doc,
        ['지표', '3개월', '6개월', '12개월'],
        [
            ['뉴스레터 구독자', '200', '1,000', '3,000'],
            ['뉴스레터 오픈율', '35%', '40%', '40%+'],
            ['상담 신청 건수', '월 5', '월 15', '월 30'],
            ['리드 매그넷 다운로드', '50', '300', '1,000'],
            ['AI 매칭 퀴즈 완료 수', '-', '500', '3,000'],
            ['상담→계약 전환', '20%', '25%', '30%'],
        ]
    )
    doc.add_paragraph()

    doc.add_heading('Freemium 전환 판단 지표', level=2)
    p = doc.add_paragraph()
    p.add_run('Phase 3 유료화 결정 시 참고할 데이터 포인트:')
    create_styled_table(doc,
        ['지표', '유료화 가능 기준', '측정 방법'],
        [
            ['DB 반복 방문율', '주 2회 이상 방문 사용자 30%+', 'GA4'],
            ['AI 매칭 사용률', '월간 활성 사용자 20%+', '내부 로그'],
            ['뉴스레터 유료 의향', '설문 응답자 15%+ "유료 지불 의향"', '설문조사'],
            ['특정 기능 집중도', '하나의 기능에 사용자 시간 40%+ 집중', 'GA4'],
        ]
    )

    doc.add_page_break()

    # ─── 16. 리스크 관리 ───
    doc.add_heading('16. 리스크 관리', level=1)

    create_styled_table(doc,
        ['리스크', '영향도', '대응 방안'],
        [
            ['초기 트래픽 부족', '높음', '정부지원 DB로 실용 트래픽 선확보, SEO 집중, GTM 전략 실행'],
            ['정부지원 데이터 수집 차단', '높음', '공공데이터포털 API 우선, 다수 소스 확보, RSS 활용'],
            ['시민기자 투고 부족', '중간', '센터 기존 고객 대상 초기 시딩, 투고 인센티브'],
            ['콘텐츠 품질 불균일', '중간', 'AI 기반 검수 보조 + 편집 가이드라인'],
            ['AI/DX 콘텐츠 진부화', '중간', '도구 트렌드 월간 업데이트, 외부 전문가 기고'],
            ['경쟁 미디어 유사 서비스', '낮음', '정부지원 DB + AI 매칭 + 커뮤니티 = 삼중 방어벽'],
            ['저작권 이슈 (스크래핑)', '중간', '요약/링크 방식, RSS 우선, 원문 출처 명시'],
            ['5대 카테고리 범위 과다', '중간', '우선순위 설정으로 리소스 집중'],
            ['Freemium 수익 지연', '중간', '컨설팅 리드로 초기 수익 확보, 6개월 후 데이터 기반 전환'],
        ]
    )

    doc.add_page_break()

    # ─── 17. 비용 구조 ───
    doc.add_heading('17. 비용 구조', level=1)

    create_styled_table(doc,
        ['항목', '월 비용', '비고'],
        [
            ['호스팅 (Railway)', '$5~20', '트래픽에 따라 스케일링'],
            ['도메인', '$1~2 (연 $12~20)', '커스텀 도메인'],
            ['뉴스레터 (Resend)', '$0~20', '구독자 수에 따라'],
            ['AI API (GPT 등)', '$10~50', '큐레이션/검수 자동화'],
            ['콘텐츠 제작', '인건비', '센터 인력 활용'],
            ['총 초기 운영비', '월 $30~90', '극히 낮은 초기 비용'],
        ],
        header_color='0D7377'
    )
    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('핵심: ')
    run.bold = True
    p.add_run('센터 인력 활용 + AI 자동화로 초기 운영 비용을 최소화하며, '
              '컨설팅 리드를 통한 수익으로 운영비를 충당합니다.')

    # ─── 마지막 페이지 ───
    doc.add_page_break()

    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('The Founders Korea')
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('창업가를 위한 실전 미디어 플랫폼')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('한국창업경영센터')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x0D, 0x73, 0x77)
    run.bold = True

    doc.add_paragraph()

    contact_info = [
        '문의: contact@founderskorea.com',
        '웹사이트: https://founderskorea.com',
    ]
    for info_text in contact_info:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(info_text)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    # ─── 저장 ───
    output_path = '/mnt/c/PycharmProjects/Founderskorea/The_Founders_Korea_사업전략서.docx'
    doc.save(output_path)
    print(f'문서 생성 완료: {output_path}')


if __name__ == '__main__':
    main()
